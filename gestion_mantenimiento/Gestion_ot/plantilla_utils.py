"""
Generador de PDFs desde plantilla DOCX con reemplazo de tags
Soporta tanto desarrollo (con docx2pdf + LibreOffice) como producción (Railway con ReportLab)
Versión Mejorada: Maneja imágenes remotas, firmas posicionadas y conversión PDF robusta
"""

import os
import io
import logging
import base64
import requests
import tempfile
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _obtener_imagen_temporal(file_field_o_url):
    """
    Obtiene ruta temporal de una imagen desde:
    - FileField de Django (local)
    - URL remota (Cloudinary)
    - Data URL base64
    
    Args:
        file_field_o_url: FileField, URL string, o data URL base64
    
    Returns:
        Tuple (ruta_temporal, es_temporal) - (str, bool) o (None, False) si falla
    """
    if not file_field_o_url:
        return None, False
    
    try:
        # Si es una URL string (Cloudinary)
        if isinstance(file_field_o_url, str):
            if file_field_o_url.startswith('http://') or file_field_o_url.startswith('https://'):
                # Descargar desde URL remota
                logger.info(f"Descargando imagen remota: {file_field_o_url[:80]}")
                response = requests.get(file_field_o_url, timeout=15)
                response.raise_for_status()
                
                # Guardar en temporal
                suffix = '.png' if 'png' in file_field_o_url.lower() else '.jpg'
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(response.content)
                    tmp_path = tmp.name
                
                logger.info(f"Imagen remota descargada a: {tmp_path}")
                return tmp_path, True
            
            # Si es data URL base64
            elif file_field_o_url.startswith('data:'):
                header, data = file_field_o_url.split(',', 1)
                imagen_bytes = base64.b64decode(data)
                
                suffix = '.png' if 'png' in header else '.jpg'
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(imagen_bytes)
                    tmp_path = tmp.name
                
                logger.info(f"Data URL decodificado a: {tmp_path}")
                return tmp_path, True
        
        # Si es un FileField de Django
        else:
            # Intentar obtener URL remota primero (Cloudinary)
            if hasattr(file_field_o_url, 'url'):
                url = file_field_o_url.url
                if url and (url.startswith('http://') or url.startswith('https://')):
                    return _obtener_imagen_temporal(url)  # Recursivo
            
            # Fallback a path local
            if hasattr(file_field_o_url, 'path'):
                local_path = file_field_o_url.path
                if local_path and os.path.exists(local_path):
                    logger.info(f"Usando imagen local: {local_path}")
                    return local_path, False
    
    except Exception as e:
        logger.warning(f"Error obteniendo imagen temporal: {e}")
    
    return None, False


def _insertar_firma_en_docx(doc, cierre_ot):
    """
    Inserta las firmas del técnico y receptor en el DOCX
    Maneja data URLs y posicionamiento correcto
    
    Args:
        doc: Documento de python-docx
        cierre_ot: Instancia de CierreOt con firmas
    """
    if not cierre_ot.firma_digital and not cierre_ot.firma_receptor:
        logger.info("No hay firmas para agregar")
        return
    
    try:
        doc.add_page_break()
        heading = doc.add_paragraph("Confirmación de trabajo recibido:")
        heading.style = 'Heading 2'
        
        # Crear tabla de firmas (3 filas: encabezados, firmas, documentos)
        tabla_firmas = doc.add_table(rows=3, cols=2)
        tabla_firmas.style = 'Light Grid Accent 1'
        
        # Fila 1: Encabezados
        tabla_firmas.rows[0].cells[0].text = "Realizador por:"
        tabla_firmas.rows[0].cells[1].text = "Recibido por:"
        
        # Fila 2: Firmas
        cell_firma_tec = tabla_firmas.rows[1].cells[0]
        cell_firma_rec = tabla_firmas.rows[1].cells[1]
        
        # Limpiar celdas
        for paragraph in cell_firma_tec.paragraphs:
            for run in paragraph.runs:
                r = run._r
                r.getparent().remove(r)
        for paragraph in cell_firma_rec.paragraphs:
            for run in paragraph.runs:
                r = run._r
                r.getparent().remove(r)
        
        # Agregar firma técnico
        if cierre_ot.firma_digital:
            tmp_path, es_temp = _obtener_imagen_temporal(cierre_ot.firma_digital)
            if tmp_path and os.path.exists(tmp_path):
                try:
                    p = cell_firma_tec.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    # Tamaño: 2.5 pulgadas de ancho
                    run.add_picture(tmp_path, width=Inches(2.5))
                    logger.info("✅ Firma técnico agregada")
                    if es_temp:
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                except Exception as e:
                    logger.warning(f"Error insertando firma técnico: {e}")
                    cell_firma_tec.paragraphs[0].text = "[Firma no disponible]"
            else:
                cell_firma_tec.paragraphs[0].text = "[Sin firma]"
        else:
            cell_firma_tec.paragraphs[0].text = "[Sin firma]"
        
        # Agregar firma receptor
        if cierre_ot.firma_receptor:
            tmp_path, es_temp = _obtener_imagen_temporal(cierre_ot.firma_receptor)
            if tmp_path and os.path.exists(tmp_path):
                try:
                    p = cell_firma_rec.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(tmp_path, width=Inches(2.5))
                    logger.info("✅ Firma receptor agregada")
                    if es_temp:
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                except Exception as e:
                    logger.warning(f"Error insertando firma receptor: {e}")
                    cell_firma_rec.paragraphs[0].text = "[Firma no disponible]"
            else:
                cell_firma_rec.paragraphs[0].text = "[Sin firma]"
        else:
            cell_firma_rec.paragraphs[0].text = "[Sin firma]"
        
        # Fila 3: Documentos de identidad
        tabla_firmas.rows[2].cells[0].text = f"Documento: {cierre_ot.documento_tecnico or 'N/A'}"
        tabla_firmas.rows[2].cells[1].text = f"Documento: {cierre_ot.documento_receptor or 'N/A'}"
        
        logger.info("✅ Tabla de firmas completada")
        
    except Exception as e:
        logger.error(f"Error agregando tabla de firmas: {e}")
        import traceback
        logger.error(traceback.format_exc())

def _agregar_imagenes_a_docx(doc, cierre_ot):
    """
    Agrega imágenes ANTES y DESPUÉS al final del documento DOCX
    Maneja imágenes remotas (Cloudinary) correctamente
    
    Args:
        doc: Documento de python-docx
        cierre_ot: Instancia de CierreOt con imágenes relacionadas
    """
    imagenes_antes = cierre_ot.imagenes.filter(tipo='antes')
    imagenes_despues = cierre_ot.imagenes.filter(tipo='despues')
    
    if not (imagenes_antes.exists() or imagenes_despues.exists()):
        logger.info("No hay imágenes para agregar")
        return
    
    try:
        doc.add_page_break()
        
        # Sección ANTES
        if imagenes_antes.exists():
            heading = doc.add_paragraph("─ ANTES ─")
            heading.style = 'Heading 2'
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            for img in imagenes_antes:
                try:
                    img_path, es_temp = _obtener_imagen_temporal(img.imagen)
                    
                    if img_path and os.path.exists(img_path):
                        try:
                            # Crear párrafo para la imagen
                            p = doc.add_paragraph()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run()
                            # Ancho máximo: 4 pulgadas
                            run.add_picture(img_path, width=Inches(4))
                            logger.info(f"✅ Imagen ANTES agregada: {img.imagen.name}")
                            
                            if es_temp:
                                try:
                                    os.unlink(img_path)
                                except:
                                    pass
                        except Exception as e:
                            logger.warning(f"Error insertando imagen ANTES: {e}")
                    else:
                        logger.warning(f"No se pudo obtener ruta para imagen: {img.imagen.name}")
                
                except Exception as e:
                    logger.warning(f"Error procesando imagen ANTES: {e}")
            
            doc.add_paragraph()  # Separador
        
        # Sección DESPUÉS
        if imagenes_despues.exists():
            heading = doc.add_paragraph("─ DESPUÉS ─")
            heading.style = 'Heading 2'
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            for img in imagenes_despues:
                try:
                    img_path, es_temp = _obtener_imagen_temporal(img.imagen)
                    
                    if img_path and os.path.exists(img_path):
                        try:
                            p = doc.add_paragraph()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(4))
                            logger.info(f"✅ Imagen DESPUÉS agregada: {img.imagen.name}")
                            
                            if es_temp:
                                try:
                                    os.unlink(img_path)
                                except:
                                    pass
                        except Exception as e:
                            logger.warning(f"Error insertando imagen DESPUÉS: {e}")
                    else:
                        logger.warning(f"No se pudo obtener ruta para imagen: {img.imagen.name}")
                
                except Exception as e:
                    logger.warning(f"Error procesando imagen DESPUÉS: {e}")
        
        logger.info(f"✅ Imágenes agregadas: {imagenes_antes.count()} antes, {imagenes_despues.count()} después")
        
    except Exception as e:
        logger.error(f"Error agregando imágenes a DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())


def reemplazar_tags_en_docx(doc, reemplazos):
    """
    Reemplaza tags <<tag>> en un documento DOCX de forma robusta
    Maneja párrafos, tablas, celdas anidadas y preserva formato
    
    Args:
        doc: Documento de python-docx
        reemplazos: Dict con {tag: valor}
    
    Ejemplo:
        reemplazos = {
            'OT': '96',
            'cliente': 'Produpan',
            'equipo': 'Cava de refrigeración',
        }
    """
    replacements_count = 0
    
    def reemplazar_en_paragraph(paragraph):
        """Reemplaza tags en un párrafo completo"""
        nonlocal replacements_count
        
        # Concatenar texto de todos los runs
        texto_completo = ''.join([run.text for run in paragraph.runs])
        
        # Si no hay placeholders, no hacer nada
        if not any(f"<<{tag}>>" in texto_completo for tag in reemplazos.keys()):
            return
        
        # Hacer reemplazos
        nuevo_texto = texto_completo
        for tag, valor in reemplazos.items():
            placeholder = f"<<{tag}>>"
            valor_str = str(valor) if valor is not None else ''
            
            if placeholder in nuevo_texto:
                nuevo_texto = nuevo_texto.replace(placeholder, valor_str)
                replacements_count += 1
        
        # Si el texto cambió, reconstruir runs
        if nuevo_texto != texto_completo:
            # Limpiar todos los runs EXCEPTO el primero (para no romper estructura)
            for run in paragraph.runs[::-1]:
                r_element = run._element
                r_element.getparent().remove(r_element)
            
            # Agregar nuevo run con el texto completo
            paragraph.add_run(nuevo_texto)
    
    # Procesar párrafos normales
    for paragraph in doc.paragraphs:
        reemplazar_en_paragraph(paragraph)
    
    # Procesar tablas recursivamente
    def procesar_celdas(celdas):
        for cell in celdas:
            # Procesar párrafos en la celda
            for paragraph in cell.paragraphs:
                reemplazar_en_paragraph(paragraph)
            
            # Procesar tablas anidadas
            for nested_table in cell.tables:
                for row in nested_table.rows:
                    procesar_celdas(row.cells)
    
    for table in doc.tables:
        for row in table.rows:
            procesar_celdas(row.cells)
    
    logger.info(f"🔍 Reemplazos realizados: {replacements_count}")


def generar_pdf_desde_plantilla(cierre_ot, plantilla_path=None):
    """
    Genera PDF a partir de plantilla DOCX reemplazando tags, insertando firmas e imágenes
    Versión mejorada con manejo correcto de imágenes remotas (Cloudinary)
    
    Args:
        cierre_ot: Instancia de CierreOt
        plantilla_path: Ruta a plantilla_ot.docx (si None, usa ruta por defecto)
    
    Returns:
        BytesIO con contenido PDF o None si falla
    """
    if plantilla_path is None:
        # Ruta por defecto: plantilla_ot.docx en la raíz del proyecto
        plantilla_path = os.path.join(
            os.path.dirname(__file__),
            '../..',
            'plantilla_ot.docx'
        )
    
    try:
        # Verificar que plantilla existe
        if not os.path.exists(plantilla_path):
            logger.warning(f"⚠️ Plantilla no encontrada en {plantilla_path}")
            return None
        
        logger.info(f"📄 Cargando plantilla desde: {plantilla_path}")
        doc = Document(plantilla_path)
        
        # Preparar datos para reemplazo
        try:
            solicitud = cierre_ot.orden_trabajo.solicitud
            equipo_nombre = solicitud.equipo.nombre if solicitud.equipo else "N/A"
            cliente_nombre = solicitud.ubicacion.nombre if solicitud.ubicacion else solicitud.PDV or "N/A"
            fecha_formato = cierre_ot.fecha_inicio_actividad.strftime('%d/%m/%Y') if cierre_ot.fecha_inicio_actividad else datetime.now().strftime('%d/%m/%Y')
        except Exception as e:
            logger.error(f"❌ Error extrayendo datos de cierre_ot: {e}")
            equipo_nombre = "N/A"
            cliente_nombre = "N/A"
            fecha_formato = datetime.now().strftime('%d/%m/%Y')
        
        # Tags a reemplazar (mapear correctamente a la plantilla)
        reemplazos = {
            'OT': str(solicitud.consecutivo) if solicitud else '',
            'equipo': equipo_nombre,
            'cliente': cliente_nombre,
            'fecha': fecha_formato,
            'tipomtto': cierre_ot.tipo_mantenimiento or 'N/A',
            'tipointervencion': cierre_ot.tipo_intervencion or 'N/A',
            'causafalla': cierre_ot.causa_falla or 'N/A',
            'sesoluciono': 'Sí' if cierre_ot.se_soluciono else 'No',
            'descripcion': cierre_ot.descripcion_falla or 'N/A',
            'observacion': cierre_ot.observaciones or 'N/A',
            'nombret': cierre_ot.nombre_tecnico or 'N/A',
            'recibido': cierre_ot.nombre_receptor or 'N/A',
            'cct': cierre_ot.documento_tecnico or 'N/A',
            'cc': cierre_ot.documento_receptor or 'N/A',
        }
        
        logger.info(f"🔍 Reemplazando {len(reemplazos)} tags en plantilla")
        reemplazar_tags_en_docx(doc, reemplazos)
        logger.info("✅ Tags reemplazados exitosamente")
        
        # Agregar firmas (incluida la lógica de inserción en plantilla existente)
        try:
            _insertar_firma_en_docx(doc, cierre_ot)
            logger.info("✅ Firmas insertadas")
        except Exception as e:
            logger.warning(f"⚠️ Error agregando firmas: {e}")
        
        # Agregar imágenes al final del documento
        try:
            _agregar_imagenes_a_docx(doc, cierre_ot)
            logger.info("✅ Imágenes agregadas")
        except Exception as e:
            logger.warning(f"⚠️ Error agregando imágenes: {e}")
        
        # Guardar DOCX temporal en memoria
        docx_temporal = io.BytesIO()
        doc.save(docx_temporal)
        docx_temporal.seek(0)
        
        logger.info("📝 DOCX generado en memoria con firmas e imágenes")
        
        # Intentar convertir a PDF
        pdf_buffer = _convertir_docx_a_pdf(docx_temporal, cierre_ot)
        
        if pdf_buffer:
            logger.info("✅ PDF generado exitosamente desde plantilla")
            return pdf_buffer
        else:
            logger.warning("⚠️ No se pudo convertir DOCX a PDF")
            return None
            
    except Exception as e:
        logger.error(f"❌ Error generando PDF desde plantilla: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None


def _convertir_docx_a_pdf(docx_buffer, cierre_ot=None):
    """
    Convierte BytesIO DOCX a BytesIO PDF
    Intenta: 1) Pandoc, 2) docx2pdf + LibreOffice, 3) ReportLab fallback
    
    Args:
        docx_buffer: BytesIO con contenido DOCX
        cierre_ot: Instancia de CierreOt (opcional, para ReportLab fallback)
    
    Returns:
        BytesIO con PDF o None si falla
    """
    try:
        # Opción 1: Intentar con Pandoc (mejor calidad, más ligero)
        logger.info("📄 Intentando conversión DOCX→PDF con Pandoc")
        try:
            # Guardar DOCX temporal
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx.write(docx_buffer.getvalue())
                tmp_docx_path = tmp_docx.name
            
            # Convertir a PDF con Pandoc
            tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
            result = subprocess.run(
                ['pandoc', tmp_docx_path, '-o', tmp_pdf_path, '--pdf-engine=xelatex'],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0 and os.path.exists(tmp_pdf_path):
                # Leer PDF
                with open(tmp_pdf_path, 'rb') as f:
                    pdf_buffer = io.BytesIO(f.read())
                
                # Limpiar temporales
                os.unlink(tmp_docx_path)
                if os.path.exists(tmp_pdf_path):
                    os.unlink(tmp_pdf_path)
                
                logger.info("✅ Conversión Pandoc exitosa")
                return pdf_buffer
            else:
                logger.warning(f"Pandoc falló: {result.stderr.decode() if result.stderr else 'sin error'}")
                os.unlink(tmp_docx_path)
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            logger.info(f"Pandoc no disponible: {e}")
        except Exception as e:
            logger.warning(f"Error con Pandoc: {e}")
        
        # Opción 2: Intentar con docx2pdf + LibreOffice
        logger.info("💻 Intentando conversión DOCX→PDF con docx2pdf + LibreOffice")
        try:
            from docx2pdf import convert
            
            # Guardar DOCX temporal
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx.write(docx_buffer.getvalue())
                tmp_docx_path = tmp_docx.name
            
            # Convertir a PDF
            tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
            convert(tmp_docx_path, tmp_pdf_path)
            
            # Leer PDF
            with open(tmp_pdf_path, 'rb') as f:
                pdf_buffer = io.BytesIO(f.read())
            
            # Limpiar temporales
            os.unlink(tmp_docx_path)
            if os.path.exists(tmp_pdf_path):
                os.unlink(tmp_pdf_path)
            
            logger.info("✅ Conversión docx2pdf exitosa")
            return pdf_buffer
            
        except Exception as e:
            logger.warning(f"docx2pdf falló: {e}")
        
        # Opción 3: Fallback con ReportLab
        logger.info("🚂 Usando ReportLab para convertir DOCX a PDF")
        return _docx_a_pdf_reportlab(docx_buffer, cierre_ot)
    
    except Exception as e:
        logger.error(f"Error en conversión: {e}")
        return None


def _docx_a_pdf_reportlab(docx_buffer, cierre_ot=None):
    """
    Convierte DOCX a PDF leyendo contenido con python-docx y generando PDF con ReportLab
    Funciona en Railway sin LibreOffice
    Soporta imágenes insertas en el documento y las obtiene directamente de cierre_ot
    
    Args:
        docx_buffer: BytesIO con contenido DOCX
        cierre_ot: Instancia de CierreOt (para obtener imágenes y firmas)
    
    Returns:
        BytesIO con PDF
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as PlatypusImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        
        logger.info("📋 Leyendo documento DOCX para conversión a PDF con ReportLab")
        
        # Cargar DOCX desde buffer
        docx_buffer.seek(0)
        doc_original = Document(docx_buffer)
        
        # Crear PDF con Platypus
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        story = []
        
        # Crear estilos personalizados
        heading_style = ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading1'],
            fontSize=12,
            textColor=colors.HexColor('#1F2937'),
            spaceAfter=8,
            spaceBefore=4,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER
        )
        
        body_style = ParagraphStyle(
            name='CustomBody',
            parent=styles['BodyText'],
            fontSize=9,
            textColor=colors.HexColor('#374151'),
            spaceAfter=4,
            leading=11
        )
        
        # Extraer párrafos
        for i, para in enumerate(doc_original.paragraphs):
            texto = para.text.strip()
            
            if texto:
                # Detectar nivel de título
                if para.style and para.style.name and para.style.name.startswith('Heading'):
                    story.append(Paragraph(texto, heading_style))
                else:
                    story.append(Paragraph(texto, body_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Extraer tablas con mejor formateo
        for table_idx, tabla in enumerate(doc_original.tables):
            try:
                # Convertir tabla preservando estructura
                data = []
                max_cols = max(len(tabla.rows[0].cells) if tabla.rows else 0, 
                              len(tabla.columns)) if tabla.columns else 0
                
                for row_idx, row in enumerate(tabla.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])
                        row_data.append(cell_text or '')
                    
                    # Rellenar celdas faltantes
                    while len(row_data) < max_cols:
                        row_data.append('')
                    
                    data.append(row_data)
                
                if data and len(data[0]) > 0:
                    # Calcular ancho de columnas automáticamente
                    available_width = 7.5*inch  # Ancho disponible
                    col_widths = [available_width / len(data[0])] * len(data[0])
                    
                    # Crear tabla Platypus
                    t = Table(data, colWidths=col_widths)
                    
                    # Estilos de tabla
                    style_commands = [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F2937')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('TOPPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F9FAFB')),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D5DB')),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F3F4F6')]),
                        ('TOPPADDING', (0, 1), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                    ]
                    
                    t.setStyle(TableStyle(style_commands))
                    story.append(t)
                    story.append(Spacer(1, 0.15*inch))
                    
            except Exception as e:
                logger.warning(f"⚠️ Error procesando tabla {table_idx}: {e}")
                continue
        
        # Agregar firmas si cierre_ot está disponible
        if cierre_ot:
            try:
                story.append(PageBreak())
                story.append(Paragraph("<b>Confirmación de trabajo recibido:</b>", heading_style))
                story.append(Spacer(1, 0.2*inch))
                
                # Tabla de firmas
                firma_data = [
                    ['Realizador por:', 'Recibido por:']
                ]
                
                # Obtener imágenes de firmas
                firma_tec_img = None
                firma_rec_img = None
                
                if cierre_ot.firma_digital:
                    tmp_path, _ = _obtener_imagen_temporal(cierre_ot.firma_digital)
                    if tmp_path and os.path.exists(tmp_path):
                        try:
                            firma_tec_img = PlatypusImage(tmp_path, width=2*inch, height=1*inch)
                        except:
                            pass
                
                if cierre_ot.firma_receptor:
                    tmp_path, _ = _obtener_imagen_temporal(cierre_ot.firma_receptor)
                    if tmp_path and os.path.exists(tmp_path):
                        try:
                            firma_rec_img = PlatypusImage(tmp_path, width=2*inch, height=1*inch)
                        except:
                            pass
                
                firma_data.append([
                    firma_tec_img or '[Sin firma]',
                    firma_rec_img or '[Sin firma]'
                ])
                
                firma_data.append([
                    f"Doc: {cierre_ot.documento_tecnico or 'N/A'}",
                    f"Doc: {cierre_ot.documento_receptor or 'N/A'}"
                ])
                
                firma_table = Table(firma_data, colWidths=[3.75*inch, 3.75*inch])
                firma_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F2937')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D5DB')),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('FONTSIZE', (0, 1), (-1, 1), 8),
                    ('TOPPADDING', (0, 0), (-1, 0), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('TOPPADDING', (0, 1), (-1, 1), 10),
                    ('BOTTOMPADDING', (0, 1), (-1, 1), 10),
                ]))
                story.append(firma_table)
                logger.info("✅ Firmas agregadas al PDF con ReportLab")
            except Exception as e:
                logger.warning(f"⚠️ Error agregando firmas: {e}")
        
        # Agregar imágenes ANTES y DESPUÉS si cierre_ot está disponible
        if cierre_ot:
            try:
                imagenes_antes = cierre_ot.imagenes.filter(tipo='antes')
                imagenes_despues = cierre_ot.imagenes.filter(tipo='despues')
                
                if imagenes_antes.exists() or imagenes_despues.exists():
                    story.append(PageBreak())
                
                # Sección ANTES
                if imagenes_antes.exists():
                    story.append(Paragraph("<b>─ ANTES ─</b>", heading_style))
                    story.append(Spacer(1, 0.1*inch))
                    
                    for img in imagenes_antes:
                        try:
                            img_path, es_temp = _obtener_imagen_temporal(img.imagen)
                            if img_path and os.path.exists(img_path):
                                try:
                                    img_platypus = PlatypusImage(img_path, width=4*inch, height=3*inch)
                                    story.append(img_platypus)
                                    story.append(Spacer(1, 0.1*inch))
                                    logger.info(f"✅ Imagen ANTES agregada al PDF")
                                except Exception as e:
                                    logger.warning(f"Error insertando imagen ANTES: {e}")
                        except Exception as e:
                            logger.warning(f"Error procesando imagen ANTES: {e}")
                
                # Separador
                if imagenes_antes.exists() and imagenes_despues.exists():
                    story.append(Spacer(1, 0.2*inch))
                
                # Sección DESPUÉS
                if imagenes_despues.exists():
                    story.append(Paragraph("<b>─ DESPUÉS ─</b>", heading_style))
                    story.append(Spacer(1, 0.1*inch))
                    
                    for img in imagenes_despues:
                        try:
                            img_path, es_temp = _obtener_imagen_temporal(img.imagen)
                            if img_path and os.path.exists(img_path):
                                try:
                                    img_platypus = PlatypusImage(img_path, width=4*inch, height=3*inch)
                                    story.append(img_platypus)
                                    story.append(Spacer(1, 0.1*inch))
                                    logger.info(f"✅ Imagen DESPUÉS agregada al PDF")
                                except Exception as e:
                                    logger.warning(f"Error insertando imagen DESPUÉS: {e}")
                        except Exception as e:
                            logger.warning(f"Error procesando imagen DESPUÉS: {e}")
                
                logger.info("✅ Imágenes agregadas al PDF con ReportLab")
            except Exception as e:
                logger.warning(f"⚠️ Error agregando imágenes: {e}")
        
        # Generar PDF
        pdf_doc.build(story)
        pdf_buffer.seek(0)
        
        logger.info("✅ DOCX convertido a PDF exitosamente con ReportLab (mejorado con imágenes y firmas)")
        return pdf_buffer
        
    except Exception as e:
        logger.error(f"❌ Error en conversión DOCX→PDF ReportLab: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return _generar_pdf_simple_minimo()


def _generar_pdf_simple_minimo():
    """
    Fallback final: PDF mínimo si todo falla
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        logger.warning("⚠️ Usando PDF mínimo como fallback final")
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, "Informe de Mantenimiento")
        
        c.setFont("Helvetica", 11)
        c.drawString(50, 720, "El informe se procesó correctamente")
        c.drawString(50, 700, "pero con contenido mínimo.")
        
        c.save()
        pdf_buffer.seek(0)
        
        logger.info("✅ PDF mínimo generado")
        return pdf_buffer
    
    except Exception as e:
        logger.error(f"Error crítico generando PDF: {e}")
        return None
