"""
Script para crear la plantilla DOCX CORRECTA para PDF de Cierre de OT
Basada en el diseño que funciona bien: logo THERMOVOLT, tablas limpias, estructura simple
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def crear_plantilla_correcta(salida_path, logo_path=None):
    """
    Crea la plantilla DOCX perfecta para Cierre de OT con logo real
    """
    doc = Document()
    
    # Márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # =====================================================
    # ENCABEZADO: LOGO
    # =====================================================
    
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if logo_path and os.path.exists(logo_path):
        try:
            run = p_logo.add_run()
            run.add_picture(logo_path, width=Inches(2.5))
            print(f"✅ Logo agregado: {logo_path}")
        except Exception as e:
            print(f"⚠️ Error al agregar logo: {e}")
            p_logo.text = "[LOGO THERMOVOLT]"
    else:
        p_logo.text = "[LOGO THERMOVOLT]"
    
    # =====================================================
    # SERVICIO N°
    # =====================================================
    
    tabla_servicio = doc.add_table(rows=1, cols=2)
    tabla_servicio.autofit = False
    
    # Celda 1: "Servicio N°"
    tabla_servicio.rows[0].cells[0].text = "Servicio\nN°"
    for para in tabla_servicio.rows[0].cells[0].paragraphs:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(10)
    
    # Celda 2: Valor
    tabla_servicio.rows[0].cells[1].text = "<<OT>>"
    for para in tabla_servicio.rows[0].cells[1].paragraphs:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(12)
    
    # Agregar espacio
    doc.add_paragraph()
    
    # =====================================================
    # TABLA 1: EQUIPO, CLIENTE, FECHA
    # =====================================================
    
    tabla1 = doc.add_table(rows=2, cols=3)
    tabla1.style = 'Light Grid Accent 1'
    
    # Encabezados
    tabla1.rows[0].cells[0].text = "Equipo"
    tabla1.rows[0].cells[1].text = "Cliente"
    tabla1.rows[0].cells[2].text = "FECHA"
    
    # Datos
    tabla1.rows[1].cells[0].text = "<<equipo>>"
    tabla1.rows[1].cells[1].text = "<<cliente>>"
    tabla1.rows[1].cells[2].text = "<<fecha>>"
    
    # Estilo de encabezados
    for cell in tabla1.rows[0].cells:
        for para in cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.bold = True
    
    # =====================================================
    # TABLA 2: TIPO DE MANTENIMIENTO
    # =====================================================
    
    tabla2 = doc.add_table(rows=2, cols=4)
    tabla2.style = 'Light Grid Accent 1'
    
    # Encabezados
    tabla2.rows[0].cells[0].text = "TIPO DE MANTENIMIENTO"
    tabla2.rows[0].cells[1].text = "TIPO DE INTERVENCIÓN"
    tabla2.rows[0].cells[2].text = "CAUSA DE LA FALLA"
    tabla2.rows[0].cells[3].text = "SE SOLUCIONO LA FALLA"
    
    # Datos
    tabla2.rows[1].cells[0].text = "<<tipomtto>>"
    tabla2.rows[1].cells[1].text = "<<tipointervencion>>"
    tabla2.rows[1].cells[2].text = "<<causafalla>>"
    tabla2.rows[1].cells[3].text = "<<sesoluciono>>"
    
    # Estilo de encabezados
    for cell in tabla2.rows[0].cells:
        for para in cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    # Estilo de datos
    for cell in tabla2.rows[1].cells:
        for para in cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # =====================================================
    # DESCRIPCIÓN DEL TRABAJO REALIZADO
    # =====================================================
    
    p_desc_titulo = doc.add_paragraph()
    p_desc_titulo.text = "DESCRIPCIÓN DEL TRABAJO REALIZADO"
    for run in p_desc_titulo.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Tabla para la descripción
    tabla_desc = doc.add_table(rows=1, cols=1)
    tabla_desc.rows[0].cells[0].text = "<<descripcion>>"
    
    # =====================================================
    # OBSERVACIONES
    # =====================================================
    
    p_obs_titulo = doc.add_paragraph()
    p_obs_titulo.text = "OBSERVACIONES"
    for run in p_obs_titulo.runs:
        run.font.bold = True
    
    # Tabla para observaciones
    tabla_obs = doc.add_table(rows=1, cols=1)
    tabla_obs.rows[0].cells[0].text = "<<observacion>>"
    
    # =====================================================
    # CONFIRMACIÓN DE TRABAJO RECIBIDO
    # =====================================================
    
    doc.add_paragraph()  # Espacio
    
    p_confirmacion = doc.add_paragraph()
    p_confirmacion.text = "Confirmación de trabajo recibido:"
    for run in p_confirmacion.runs:
        run.font.bold = True
        run.font.size = Pt(11)
    
    # Tabla de firmas
    tabla_firmas = doc.add_table(rows=3, cols=2)
    tabla_firmas.style = 'Light Grid Accent 1'
    
    # Fila 1: Encabezados
    tabla_firmas.rows[0].cells[0].text = "Realizador por:"
    tabla_firmas.rows[0].cells[1].text = "Recibido por:"
    
    for cell in tabla_firmas.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    # Fila 2: Nombres
    tabla_firmas.rows[1].cells[0].text = "<<nombret>>"
    tabla_firmas.rows[1].cells[1].text = "<<recibido>>"
    
    # Fila 3: Documentos
    tabla_firmas.rows[2].cells[0].text = "Documento de identidad: <<cct>>"
    tabla_firmas.rows[2].cells[1].text = "Documento de identidad: <<cc>>"
    
    # Guardar plantilla
    doc.save(salida_path)
    print(f"✅ Plantilla creada en: {salida_path}")


if __name__ == "__main__":
    # Obtener ruta base del proyecto
    base_path = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    
    # Ruta del logo (está en static/images/)
    logo_path = os.path.join(base_path, 'static', 'images', 'Logoinforme.jpg')
    
    print(f"🔍 Base path: {base_path}")
    print(f"🔍 Buscando logo en: {logo_path}")
    if os.path.exists(logo_path):
        print(f"✅ Logo encontrado!")
    else:
        print(f"⚠️ Logo NO encontrado, se usará placeholder")
    
    # Crear plantilla en raíz
    ruta_raiz = os.path.join(
        os.path.dirname(__file__),
        '../..',
        'plantilla_ot.docx'
    )
    
    crear_plantilla_correcta(ruta_raiz, logo_path)
    
    # También en static
    ruta_static = os.path.join(
        os.path.dirname(__file__),
        '../static',
        'plantilla_ot.docx'
    )
    
    crear_plantilla_correcta(ruta_static, logo_path)
