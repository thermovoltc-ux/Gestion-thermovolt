from docx import Document

doc = Document('plantilla_ot.docx')
print(f"=== ESTRUCTURA DE PLANTILLA ===")
print(f"Párrafos: {len(doc.paragraphs)}")
print(f"Tablas: {len(doc.tables)}")

# Contar imágenes
img_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.target_ref:
        img_count += 1
        print(f"  Imagen: {rel.target_ref}")

print(f"Total imágenes: {img_count}")

# Primeros párrafos
print(f"\n=== PRIMEROS PÁRRAFOS ===")
for i, p in enumerate(doc.paragraphs[:10]):
    if p.text.strip():
        print(f"  {i}: {p.text[:70]}")

# Tablas
print(f"\n=== TABLAS ===")
for i, table in enumerate(doc.tables):
    print(f"  Tabla {i}: {len(table.rows)} filas x {len(table.columns)} columnas")
    if table.rows:
        first_row = [c.text[:25] for c in table.rows[0].cells]
        print(f"    Encabezado: {first_row}")
